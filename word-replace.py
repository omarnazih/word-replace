from docx import Document
import os


def main():
    template_file_path = 'tagdeed_elmror.doc'
    output_file_path = 'result.doc'
    
    variables = {
        "{TRNS_DATE}": "10/12/2023",
        "{TRAFFIC_NAME}": "Test TRAFFIC_NAME",
        "{RVWD}": "Test RVWD",
        "{CMPNY_NAME_AR}": "Test CMPNY_NAME_AR",
        "{TAFWEED}": "Test TAFWEED",
        "{ASSET_DESCR1}": "example@example.com",
        "{ASSET_DT1}": "03 Jan, 2021",
        "{PERIOD}": "PERIOD test",
        "{CMPNY_NAME_AR}": "تست",
        "{EXPIRY}": "7,000",
        "{CNTRCT_NO}": "123",
        "{TABLE_DATA}": [  # Updated to match the key in the template
            ["Row 1, Cell 1", "Row 1, Cell 2", "Row 1, Cell 3"],
            ["Row 2, Cell 1", "Row 2, Cell 2", "Row 2, Cell 3"],
            ["Row 3, Cell 1", "Row 3, Cell 2", "Row 3, Cell 3"],
        ],
    }

    try:        
        template_document = Document(template_file_path)
    except OSError:
        print("Could not open/read file:", template_file_path)        

    for variable_key, variable_value in variables.items():        
        if variable_key.endswith("_DATA}"):
            print("_DATA}")
            print(variable_key)
            # Handle table data
            if isinstance(variable_value, list):
                table = template_document.add_table(rows=1, cols=len(variable_value[0]), style="Table Grid")
                for row_data in variable_value:
                    row = table.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        row[i].text = cell_data
        else:
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)

def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                if not isinstance(value, list):
                    item.text = item.text.replace(key, value)

if __name__ == '__main__':
    main()
