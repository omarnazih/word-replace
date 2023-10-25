from docx import Document
import os

def replace_placeholders_in_docx(template_path, output_path, variables):
    document = Document(template_path)

    for paragraph in document.paragraphs:
        for key, value in variables.items():
            if key.endswith("_DATA}"):
                # Handle table replacement
                if paragraph.text == key:
                    paragraph.text = ""
                    table_data = value
                    for row in table_data:
                        cells = paragraph.add_run(row).add_break()
                else:
                    continue
            else:
                # Replace regular placeholders in the document
                for run in paragraph.runs:
                    run.text = run.text.replace(key, value)

    document.save(output_path)

if __name__ == "__main__":
    template_path = "tagdeed_elmror.doc"  # Replace with the path to your .docx template
    output_path = "result.doc"  # Path for the output document
    variables = {
        "{TRNS_DATE}": "10/12/2023",
        "{TRAFFIC_NAME}": "Test TRAFFIC_NAME",
        "{RVWD}": "Test RVWD",
        "{CMPNY_NAME_AR}": "Test CMPNY_NAME_AR",
        "{TAFWEED}": "Test TAFWEED",
        "{ASSET_DESCR1}": "example@example.com",
        "{ASSET_DT1}": "03 Jan, 2021",
        "{PERIOD}": "PERIOD test",
        "{CMPNY_NAME_AR}": "تست",
        "{EXPIRY}": "7,000",
        "{CNTRCT_NO}#": "123",
        "{TABLE_DATA}": [  # Updated to match the key in the template
            ["Row 1, Cell 1", "Row 1, Cell 2", "Row 1, Cell 3"],
            ["Row 2, Cell 1", "Row 2, Cell 2", "Row 2, Cell 3"],
            ["Row 3, Cell 1", "Row 3, Cell 2", "Row 3, Cell 3"],
        ],        
    }

    replace_placeholders_in_docx(template_path, output_path, variables)
    print(f"Modified .docx file saved at: {output_path}")
